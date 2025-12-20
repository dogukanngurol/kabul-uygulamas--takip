import streamlit as st
import sqlite3
import pandas as pd
from datetime import datetime
import hashlib
import io
import json
from docx import Document
from docx.shared import Inches

# --- 1. VERİTABANI VE YAPI ---
def init_db():
    conn = sqlite3.connect('isletme_saha_v12.db', check_same_thread=False)
    c = conn.cursor()
    c.execute('CREATE TABLE IF NOT EXISTS users (email TEXT PRIMARY KEY, password TEXT, role TEXT, name TEXT, title TEXT)')
    # Çoklu fotoğraf için 'photos_json' sütunu kullanacağız
    c.execute('''CREATE TABLE IF NOT EXISTS tasks 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, assigned_to TEXT, title TEXT, 
                  description TEXT, status TEXT, report TEXT, photos_json TEXT, updated_at TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS inventory 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, item_name TEXT, 
                  assigned_to TEXT, quantity INTEGER, updated_by TEXT)''')
    
    pw = hashlib.sha256("1234".encode()).hexdigest()
    c.execute("INSERT OR IGNORE INTO users VALUES ('admin@sirket.com', ?, 'admin', 'Genel Müdür', 'Yönetici')", (pw,))
    c.execute("INSERT OR IGNORE INTO users VALUES ('deneme123@dev.com', ?, 'worker', 'Deneme Çalışan', 'Saha Ekibi')", (pw,))
    conn.commit()
    return conn

conn = init_db()

def make_hash(p): return hashlib.sha256(str.encode(p)).hexdigest()

# --- 2. ÖZEL WORD RAPORU OLUŞTURUCU ---
def create_single_task_report(row):
    doc = Document()
    doc.add_heading('İŞ BİTİRME VE SAHA RAPORU', 0)
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'İş Başlığı:'
    hdr_cells[1].text = str(row['title'])
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Sorumlu Personel:'
    row_cells[1].text = f"{row['assigned_to']}"
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Tamamlanma Tarihi:'
    row_cells[1].text = str(row['updated_at'])

    doc.add_heading('Personel Notu / Raporu', level=2)
    doc.add_paragraph(str(row['report']))

    if row['photos_json']:
        doc.add_heading('Saha Fotoğrafları', level=2)
        photos = json.loads(row['photos_json'])
        for idx, p_hex in enumerate(photos):
            try:
                img_data = bytes.fromhex(p_hex)
                doc.add_picture(io.BytesIO(img_data), width=Inches(5))
                doc.add_paragraph(f"Fotoğraf {idx+1}")
            except:
                doc.add_paragraph(f"Hata: Fotoğraf {idx+1} yüklenemedi.")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. ANA ARAYÜZ ---
st.set_page_config(page_title="Saha Yönetim v12", layout="wide")

if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False

if not st.session_state['logged_in']:
    st.title("🔐 Şirket Giriş Paneli")
    e = st.text_input("E-posta")
    p = st.text_input("Şifre", type='password')
    if st.button("Sisteme Giriş"):
        c = conn.cursor()
        c.execute("SELECT * FROM users WHERE email=? AND password=?", (e, make_hash(p)))
        u = c.fetchone()
        if u:
            st.session_state.update({'logged_in':True, 'user_email':u[0], 'role':u[2], 'user_name':u[3], 'user_title':u[4]})
            st.rerun()
        else: st.error("Hatalı giriş!")
else:
    st.sidebar.title(f"👋 {st.session_state['user_name']}")
    st.sidebar.caption(f"📌 {st.session_state['user_title']}")
    
    if st.session_state['role'] == 'admin':
        menu = ["Ana Sayfa", "İş Atama & Takip", "Tamamlanan İşler", "Zimmet/Envanter", "Kullanıcı Yönetimi"]
    else:
        menu = ["Üstüme Atanan İşler", "Tamamlanan İşlerim", "Zimmetim"]
    
    choice = st.sidebar.radio("Menü", menu)
    if st.sidebar.button("Güvenli Çıkış"):
        st.session_state['logged_in'] = False
        st.rerun()

    # --- EKRANLAR ---
    if choice == "Ana Sayfa":
        st.header("📊 Genel Operasyon Özeti")
        tasks = pd.read_sql("SELECT status FROM tasks", conn)
        c1, c2 = st.columns(2)
        c1.metric("📌 Bekleyen İşler", len(tasks[tasks['status']=='Bekliyor']))
        c2.metric("✅ Tamamlanan İşler", len(tasks[tasks['status']=='Tamamlandı']))

    elif choice == "İş Atama & Takip":
        st.header("🆕 Yeni İş Ataması Yap")
        workers = pd.read_sql("SELECT email, name FROM users WHERE role='worker'", conn)
        with st.form("is_ata"):
            tit = st.text_input("İş Başlığı")
            who = st.selectbox("Görevli Personel", workers['email'])
            dsc = st.text_area("İş Detayı / Adres")
            if st.form_submit_button("Görevi Personel Ekranına Gönder"):
                conn.execute("INSERT INTO tasks (assigned_to, title, description, status) VALUES (?,?,?,?)", (who, tit, dsc, 'Bekliyor'))
                conn.commit()
                st.success(f"{tit} işi başarıyla atandı.")

    elif choice == "Tamamlanan İşler":
        st.header("📑 Tamamlanan İşler ve Raporlar")
        df_d = pd.read_sql("SELECT * FROM tasks WHERE status='Tamamlandı' ORDER BY updated_at DESC", conn)
        if df_d.empty:
            st.info("Henüz tamamlanan bir iş bulunmuyor.")
        else:
            for idx, row in df_d.iterrows():
                with st.expander(f"📍 {row['title']} - (Personel: {row['assigned_to']})"):
                    col_info, col_btn = st.columns([3, 1])
                    col_info.write(f"**Tamamlanma:** {row['updated_at']}")
                    col_info.write(f"**Rapor:** {row['report']}")
                    
                    # Her iş için ayrı Word indirme butonu
                    report_data = create_single_task_report(row)
                    col_btn.download_button(
                        label="📄 Word Raporu İndir",
                        data=report_data,
                        file_name=f"Rapor_{row['title']}_{row['id']}.docx",
                        key=f"dl_{row['id']}"
                    )
                    
                    # Fotoğrafları önizle
                    if row['photos_json']:
                        photos = json.loads(row['photos_json'])
                        cols = st.columns(3)
                        for i, p_hex in enumerate(photos):
                            cols[i % 3].image(bytes.fromhex(p_hex), use_container_width=True)

    elif choice == "Kullanıcı Yönetimi":
        st.header("👥 Personel ve Unvan Yönetimi")
        with st.expander("➕ Yeni Kullanıcı Tanımla"):
            with st.form("u_add"):
                n_e = st.text_input("E-posta")
                n_n = st.text_input("Ad Soyad")
                n_t = st.selectbox("Unvan", ["Müdür", "Müdür Yrd.", "Tekniker", "Saha Ekibi", "Ofis"])
                n_p = st.text_input("Şifre", type='password')
                n_r = st.selectbox("Sistem Yetkisi", ["worker", "admin"])
                if st.form_submit_button("Personeli Kaydet"):
                    conn.execute("INSERT OR IGNORE INTO users VALUES (?,?,?,?,?)", (n_e, make_hash(n_p), n_r, n_n, n_t))
                    conn.commit()
                    st.rerun()
        
        u_list = pd.read_sql("SELECT name, email, title, role FROM users", conn)
        st.table(u_list)
        for _, r in u_list.iterrows():
            if r['email'] != 'admin@sirket.com':
                if st.button(f"🗑️ {r['name']} Hesabını Sil", key=r['email']):
                    conn.execute("DELETE FROM users WHERE email=?", (r['email'],))
                    conn.commit()
                    st.rerun()

    elif choice == "Üstüme Atanan İşler":
        st.header("⏳ Tamamlanacak İşlerim")
        my_tasks = pd.read_sql(f"SELECT * FROM tasks WHERE assigned_to='{st.session_state['user_email']}' AND status='Bekliyor'", conn)
        if my_tasks.empty: st.info("Şu an bekleyen bir göreviniz bulunmuyor.")
        for _, row in my_tasks.iterrows():
            with st.expander(f"📌 {row['title']}"):
                st.write(f"**İş Detayı:** {row['description']}")
                rep = st.text_area("İş Sonu Rapor Notu", key=f"r_{row['id']}")
                # ÇOKLU FOTOĞRAF YÜKLEME
                uploaded_files = st.file_uploader("Fotoğrafları Yükle (Birden fazla seçebilirsiniz)", 
                                                 type=['jpg','png','jpeg'], 
                                                 accept_multiple_files=True,
                                                 key=f"f_{row['id']}")
                
                if st.button("İşi Tamamla ve Fotoğrafları Gönder", key=f"b_{row['id']}"):
                    if uploaded_files:
                        photo_list = []
                        for f in uploaded_files:
                            photo_list.append(f.read().hex()) # Veritabanına kaydetmek için hex formatına çeviriyoruz
                        
                        photos_json = json.dumps(photo_list)
                        conn.execute("UPDATE tasks SET status='Tamamlandı', report=?, photos_json=?, updated_at=? WHERE id=?", 
                                     (rep, photos_json, datetime.now().strftime("%d/%m/%Y %H:%M"), row['id']))
                        conn.commit()
                        st.success("Tüm fotoğraflar başarıyla yüklendi ve iş tamamlandı!")
                        st.rerun()
                    else: st.error("Lütfen en az bir fotoğraf yükleyin!")

    elif choice == "Tamamlanan İşlerim":
        st.header("✅ Geçmiş İş Kayıtlarım")
        df_history = pd.read_sql(f"SELECT title as 'İş Başlığı', report as 'Not', updated_at as 'Tarih' FROM tasks WHERE assigned_to='{st.session_state['user_email']}' AND status='Tamamlandı' ORDER BY updated_at DESC", conn)
        st.dataframe(df_history, use_container_width=True)

    elif choice == "Zimmetim" or choice == "Zimmet/Envanter":
        st.header("📦 Zimmet & Envanter")
        if st.session_state['role'] == 'admin':
            df_i = pd.read_sql("SELECT item_name as 'Eşya', assigned_to as 'Personel', quantity as 'Adet', updated_by as 'Ekleyen' FROM inventory", conn)
            st.dataframe(df_i, use_container_width=True)
        else:
            df_i = pd.read_sql(f"SELECT item_name, quantity FROM inventory WHERE assigned_to='{st.session_state['user_email']}'", conn)
            st.table(df_i)
        
        with st.form("inv_add"):
            i_n = st.text_input("Eşya/Ekipman Adı")
            i_q = st.number_input("Adet", 1)
            target = st.session_state['user_email'] if st.session_state['role'] == 'worker' else st.text_input("Personel E-postası")
            if st.form_submit_button("Envantere İşle"):
                conn.execute("INSERT INTO inventory (item_name, assigned_to, quantity, updated_by) VALUES (?,?,?,?)", 
                             (i_n, target, i_q, st.session_state['user_name']))
                conn.commit()
                st.rerun()
